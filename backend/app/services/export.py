"""
文档导出服务
支持 Word、PDF、Markdown 导出
"""
import os
import tempfile
from pathlib import Path
from typing import Optional
from datetime import datetime

from fastapi import HTTPException
from fastapi.responses import FileResponse


class ExportService:
    """文档导出服务"""
    
    def __init__(self):
        self.temp_dir = Path(tempfile.gettempdir()) / "gov_writing_exports"
        self.temp_dir.mkdir(exist_ok=True)
    
    def _sanitize_filename(self, filename: str) -> str:
        """清理文件名"""
        # 移除不合法字符
        invalid_chars = '\/:*?"<>|'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        return filename
    
    def export_to_markdown(
        self,
        title: str,
        content: str,
        metadata: dict = None
    ) -> Path:
        """
        导出为 Markdown
        
        Args:
            title: 文档标题
            content: 文档内容
            metadata: 元数据
            
        Returns:
            文件路径
        """
        filename = self._sanitize_filename(f"{title}.md")
        filepath = self.temp_dir / filename
        
        # 构建 frontmatter
        frontmatter = f"""---
title: {title}
created_at: {datetime.now().isoformat()}
---

"""
        
        full_content = frontmatter + content
        filepath.write_text(full_content, encoding="utf-8")
        
        return filepath
    
    def export_to_word(
        self,
        title: str,
        content: str,
        metadata: dict = None
    ) -> Path:
        """
        导出为 Word 文档 (.docx)
        
        Args:
            title: 文档标题
            content: 文档内容
            metadata: 元数据
            
        Returns:
            文件路径
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            # 如果没有 python-docx，生成 HTML 格式的 doc
            return self._export_to_html_doc(title, content)
        
        filename = self._sanitize_filename(f"{title}.docx")
        filepath = self.temp_dir / filename
        
        doc = Document()
        
        # 添加标题
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加内容（简单处理）
        paragraphs = content.split("\n\n")
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            # 检测标题
            if para.startswith("# "):
                doc.add_heading(para.replace("# ", ""), level=1)
            elif para.startswith("## "):
                doc.add_heading(para.replace("## ", ""), level=2)
            elif para.startswith("### "):
                doc.add_heading(para.replace("### ", ""), level=3)
            else:
                # 普通段落
                p = doc.add_paragraph(para)
                p.paragraph_format.first_line_indent = Inches(0.5)
        
        doc.save(str(filepath))
        return filepath
    
    def _export_to_html_doc(self, title: str, content: str) -> Path:
        """生成 HTML 格式的 Word 文档（备用方案）"""
        filename = self._sanitize_filename(f"{title}.doc")
        filepath = self.temp_dir / filename
        
        # 将 Markdown 简单转换为 HTML
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{ font-family: "SimSun", serif; line-height: 1.8; max-width: 800px; margin: 40px auto; padding: 20px; }}
        h1 {{ text-align: center; font-size: 22pt; margin-bottom: 30px; }}
        h2 {{ font-size: 16pt; margin-top: 24px; }}
        h3 {{ font-size: 14pt; margin-top: 18px; }}
        p {{ text-indent: 2em; margin: 12px 0; font-size: 12pt; }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <div class="content">
"""
        
        # 简单转换
        paragraphs = content.split("\n\n")
        for para in paragraphs:
            para = para.strip()
            if para.startswith("# "):
                html_content += f"    <h1>{para[2:]}</h1>\n"
            elif para.startswith("## "):
                html_content += f"    <h2>{para[3:]}</h2>\n"
            elif para.startswith("### "):
                html_content += f"    <h3>{para[4:]}</h3>\n"
            else:
                html_content += f"    <p>{para}</p>\n"
        
        html_content += """    </div>
</body>
</html>"""
        
        filepath.write_text(html_content, encoding="utf-8")
        return filepath
    
    def export_to_pdf(
        self,
        title: str,
        content: str,
        metadata: dict = None
    ) -> Path:
        """
        导出为 PDF
        
        注：PDF 导出需要额外依赖，这里先返回 HTML
        后续可以集成 WeasyPrint 或 reportlab
        
        Args:
            title: 文档标题
            content: 文档内容
            metadata: 元数据
            
        Returns:
            文件路径
        """
        # 目前返回 HTML，用户可以用浏览器打印为 PDF
        filename = self._sanitize_filename(f"{title}.pdf.html")
        filepath = self.temp_dir / filename
        
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        @media print {{
            @page {{ margin: 2.5cm; size: A4; }}
        }}
        body {{ font-family: "SimSun", "宋体", serif; line-height: 1.8; max-width: 800px; margin: 40px auto; padding: 20px; }}
        h1 {{ text-align: center; font-size: 22pt; margin-bottom: 30px; font-weight: bold; }}
        h2 {{ font-size: 16pt; margin-top: 24px; font-weight: bold; }}
        h3 {{ font-size: 14pt; margin-top: 18px; font-weight: bold; }}
        p {{ text-indent: 2em; margin: 12px 0; font-size: 12pt; text-align: justify; }}
        .print-btn {{ position: fixed; top: 20px; right: 20px; padding: 10px 20px; background: #409EFF; color: white; border: none; border-radius: 4px; cursor: pointer; }}
        @media print {{ .print-btn {{ display: none; }} }}
    </style>
</head>
<body>
    <button class="print-btn" onclick="window.print()">打印 / 另存为 PDF</button>
    <h1>{title}</h1>
    <div class="content">
"""
        
        paragraphs = content.split("\n\n")
        for para in paragraphs:
            para = para.strip()
            if para.startswith("# "):
                html_content += f"        <h1>{para[2:]}</h1>\n"
            elif para.startswith("## "):
                html_content += f"        <h2>{para[3:]}</h2>\n"
            elif para.startswith("### "):
                html_content += f"        <h3>{para[4:]}</h3>\n"
            else:
                html_content += f"        <p>{para}</p>\n"
        
        html_content += """    </div>
</body>
</html>"""
        
        filepath.write_text(html_content, encoding="utf-8")
        return filepath
    
    def get_export_response(
        self,
        title: str,
        content: str,
        format: str = "markdown"
    ):
        """
        获取导出文件响应
        
        Args:
            title: 文档标题
            content: 文档内容
            format: 导出格式 (markdown, word, pdf)
            
        Returns:
            FileResponse
        """
        format = format.lower()
        
        if format == "markdown" or format == "md":
            filepath = self.export_to_markdown(title, content)
            media_type = "text/markdown"
            filename = f"{title}.md"
            
        elif format == "word" or format == "docx":
            filepath = self.export_to_word(title, content)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            filename = f"{title}.docx"
            
        elif format == "pdf":
            filepath = self.export_to_pdf(title, content)
            media_type = "text/html"
            filename = f"{title}.pdf.html"
            
        else:
            raise HTTPException(status_code=400, detail=f"不支持的格式: {format}")
        
        if not filepath.exists():
            raise HTTPException(status_code=500, detail="导出失败")
        
        return FileResponse(
            path=str(filepath),
            filename=filename,
            media_type=media_type
        )


# 全局导出服务实例
export_service = ExportService()